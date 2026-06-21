from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "manuales"
OUTPUT.mkdir(exist_ok=True)

INK = RGBColor(23, 33, 43)
OCEAN = RGBColor(34, 87, 122)
COPPER = RGBColor(180, 95, 60)
MUTED = RGBColor(91, 104, 114)
LIGHT = "EEF3F5"
PAPER = "F7F4EC"
WHITE = RGBColor(255, 255, 255)
SITE_URL = "https://carlos-lopez-academic.vercel.app"
DATE_TEXT = "21 de junio de 2026"


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure_document(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in [
        ("Heading 1", 17, OCEAN, 18, 8),
        ("Heading 2", 14, OCEAN, 14, 6),
        ("Heading 3", 12, INK, 10, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_p.add_run(running_title)
    set_font(header_run, size=9, bold=True, color=MUTED)

    footer = section.footer
    page_number(footer.paragraphs[0])


def add_cover(doc, title, subtitle, audience):
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("PERFIL ACADÉMICO DIGITAL"), size=10, bold=True, color=COPPER)
    kicker.paragraph_format.space_after = Pt(18)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(8)
    set_font(title_p.add_run(title), size=28, bold=True, color=INK)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(26)
    set_font(subtitle_p.add_run(subtitle), size=14, color=OCEAN)

    box = doc.add_table(rows=3, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(5.6)
    for cell in box.column_cells(0):
        cell.width = Inches(5.6)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=125, bottom=125, start=180, end=180)
        shade_cell(cell, LIGHT)

    entries = [
        ("Sitio", SITE_URL),
        ("Dirigido a", audience),
        ("Versión", DATE_TEXT),
    ]
    for cell, (label, value) in zip(box.column_cells(0), entries):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(f"{label}: "), size=10, bold=True, color=INK)
        set_font(p.add_run(value), size=10, color=MUTED)

    doc.add_paragraph()
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(author.add_run("Carlos Fernando López Rengifo"), size=12, bold=True, color=INK)
    set_font(author.add_run("\nUniversidad Nacional del Centro del Perú"), size=10, color=MUTED)
    doc.add_page_break()


def add_intro(doc, purpose, scope):
    doc.add_heading("1. Propósito del manual", level=1)
    doc.add_paragraph(purpose)
    note_box(doc, "Alcance", scope)

    doc.add_heading("2. Acceso rápido", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.65), Inches(4.95)]
    headers = ["Elemento", "Información"]
    for index, cell in enumerate(table.rows[0].cells):
        cell.width = widths[index]
        shade_cell(cell, "DDE9EE")
        set_cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(headers[index]), size=10, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    for label, value in [
        ("Dirección", SITE_URL),
        ("Lectura", "No requiere iniciar sesión."),
        ("Participación", "Requiere una cuenta creada por el administrador."),
        ("Dispositivos", "Computadora, tableta o teléfono con navegador actualizado."),
    ]:
        cells = table.add_row().cells
        for index, cell in enumerate(cells):
            cell.width = widths[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(label), size=10, bold=True, color=INK)
        set_font(cells[1].paragraphs[0].add_run(value), size=10, color=INK)


def note_box(doc, label, text, caution=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.55)
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    shade_cell(cell, "FBEDE8" if caution else PAPER)
    p = cell.paragraphs[0]
    set_font(p.add_run(f"{label}. "), size=10.5, bold=True, color=COPPER if caution else OCEAN)
    set_font(p.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_steps(doc, title, steps):
    doc.add_heading(title, level=2)
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        set_font(p.add_run(step), size=11, color=INK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(item), size=11, color=INK)


def build_user_manual():
    doc = Document()
    configure_document(doc, "Manual de usuario | Perfil académico")
    add_cover(
        doc,
        "Manual de usuario",
        "Consulta de cursos, participación en foros y gestión de cuenta",
        "Estudiantes, egresados y visitantes autorizados",
    )
    add_intro(
        doc,
        "Este manual explica cómo consultar el perfil académico, revisar cursos y materiales, iniciar sesión y participar en los foros.",
        "El contenido público puede consultarse sin cuenta. Para responder, comentar o dar like es necesario usar una cuenta entregada por el docente.",
    )

    doc.add_heading("3. Navegación del sitio", level=1)
    add_bullets(
        doc,
        [
            "Perfil: presenta la información profesional y académica del docente.",
            "Trayectoria: muestra los principales hitos profesionales en orden cronológico.",
            "Cursos: reúne cursos actuales e históricos, con sus archivos disponibles.",
            "Foros: contiene conversaciones y consultas organizadas por curso.",
            "Publicaciones: presenta artículos, ponencias, informes y trabajos científicos.",
        ],
    )
    doc.add_heading("Navegación desde un teléfono", level=2)
    add_steps(
        doc,
        "Uso del menú móvil",
        [
            "Pulse el icono de menú ubicado en la esquina superior derecha.",
            "Seleccione la sección que desea consultar.",
            "Use el mismo menú para iniciar sesión, contactar al docente o abrir la administración.",
            "Pulse la X o seleccione una opción para cerrar el menú.",
        ],
    )

    doc.add_heading("4. Cuenta e inicio de sesión", level=1)
    add_steps(
        doc,
        "Ingresar con una cuenta asignada",
        [
            "Pulse Ingresar en la cabecera o en el menú móvil.",
            "Escriba el correo registrado por el docente.",
            "Introduzca la contraseña temporal recibida.",
            "Pulse Ingresar.",
            "Cuando el sistema lo solicite, establezca una contraseña personal de al menos ocho caracteres.",
        ],
    )
    note_box(
        doc,
        "Seguridad",
        "No comparta su contraseña. Si olvida el acceso, solicite al administrador que revise o regenere su cuenta.",
        caution=True,
    )

    doc.add_heading("5. Consulta de cursos y archivos", level=1)
    add_steps(
        doc,
        "Abrir materiales de un curso",
        [
            "Entre en la sección Cursos.",
            "Identifique el curso por nombre, código y periodo académico.",
            "Pulse Ver materiales.",
            "Seleccione el archivo disponible para abrirlo o descargarlo.",
        ],
    )
    add_bullets(
        doc,
        [
            "Los cursos actuales corresponden al periodo vigente.",
            "Los cursos históricos conservan materiales de periodos anteriores.",
            "Los archivos privados solo están disponibles para la administración.",
        ],
    )

    doc.add_heading("6. Participación en foros", level=1)
    add_steps(
        doc,
        "Responder en un foro",
        [
            "Inicie sesión con su cuenta.",
            "Busque el foro correspondiente al curso.",
            "Seleccione Responder al foro o elija una respuesta existente.",
            "Escriba un aporte claro y relacionado con el tema.",
            "Pulse Publicar respuesta y espere a que termine el indicador de carga.",
        ],
    )
    add_steps(
        doc,
        "Dar like",
        [
            "Inicie sesión.",
            "Pulse el icono de corazón del foro o de una respuesta.",
            "El sistema registra un solo like por cuenta para cada elemento.",
        ],
    )
    note_box(
        doc,
        "Convivencia",
        "Use un lenguaje académico y respetuoso. Evite datos personales, mensajes repetidos y contenido ajeno al curso.",
    )

    doc.add_heading("7. Cambio de contraseña y cierre de sesión", level=1)
    add_steps(
        doc,
        "Cambiar la contraseña",
        [
            "Abra el menú de usuario y pulse Cambiar contraseña.",
            "Escriba la nueva contraseña y repítala.",
            "Pulse Guardar nueva contraseña.",
        ],
    )
    add_steps(
        doc,
        "Cerrar sesión",
        [
            "Pulse Cerrar sesión en la cabecera o en el menú móvil.",
            "Compruebe que vuelva a mostrarse la opción Ingresar.",
        ],
    )

    doc.add_heading("8. Contador de visitas", level=1)
    doc.add_paragraph(
        "En el pie de página se muestra un contador discreto de visitas. El sistema registra una visita por sesión del navegador y no guarda nombres, correos ni direcciones IP en ese contador."
    )

    doc.add_heading("9. Solución de problemas", level=1)
    troubleshooting = [
        ("No puedo iniciar sesión", "Verifique correo y contraseña. Respete mayúsculas y símbolos."),
        ("No puedo responder", "Compruebe que inició sesión y que el foro está abierto."),
        ("No veo un archivo", "El archivo puede ser privado o haber sido retirado por el docente."),
        ("La página no cambia", "Espere a que desaparezca el loader y luego actualice el navegador."),
        ("Uso teléfono", "Abra el menú hamburguesa para acceder a todas las opciones."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.05), Inches(4.55)]
    for index, text in enumerate(["Situación", "Acción recomendada"]):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        shade_cell(cell, "DDE9EE")
        set_cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(text), size=10, bold=True)
    set_repeat_table_header(table.rows[0])
    for issue, action in troubleshooting:
        cells = table.add_row().cells
        for index, cell in enumerate(cells):
            cell.width = widths[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(issue), size=10, bold=True)
        set_font(cells[1].paragraphs[0].add_run(action), size=10)

    path = OUTPUT / "Manual_de_usuario_Carlos_Lopez.docx"
    doc.save(path)
    return path


def build_admin_manual():
    doc = Document()
    configure_document(doc, "Manual de administrador | Perfil académico")
    add_cover(
        doc,
        "Manual de administrador",
        "Gestión de perfil, usuarios, cursos, archivos, foros y publicaciones",
        "Carlos Fernando López Rengifo y administradores autorizados",
    )
    add_intro(
        doc,
        "Este manual describe las tareas de mantenimiento disponibles en el panel administrativo del sitio académico.",
        "El panel permite actualizar contenido y crear cuentas. Solo los usuarios registrados en la tabla de administradores de Supabase tienen permiso para ejecutar estas acciones.",
    )

    doc.add_heading("3. Acceso al panel", level=1)
    add_steps(
        doc,
        "Iniciar sesión como administrador",
        [
            f"Abra {SITE_URL}/admin.",
            "Escriba el correo administrativo.",
            "Introduzca la contraseña y pulse Ingresar.",
            "Compruebe que aparezca el encabezado Panel administrativo.",
        ],
    )
    note_box(
        doc,
        "Acceso restringido",
        "Una cuenta común puede participar en foros, pero no puede editar el sitio ni crear usuarios.",
        caution=True,
    )

    doc.add_heading("4. Gestión de usuarios", level=1)
    add_steps(
        doc,
        "Preparar el archivo",
        [
            "Descargue la plantilla desde la sección Usuarios autorizados.",
            "Complete una fila por persona.",
            "Use exactamente las columnas Nombre, Apellido y Correo.",
            "Guarde el archivo como .xlsx o .csv.",
        ],
    )
    add_steps(
        doc,
        "Importar y crear cuentas",
        [
            "Pulse Seleccionar archivo Excel o CSV.",
            "Revise la vista previa y corrija las filas marcadas con error.",
            "Pulse Crear usuarios.",
            "Espere a que finalice el loader.",
            "Descargue inmediatamente el archivo de credenciales temporales.",
            "Entregue a cada persona únicamente su propio correo y contraseña.",
        ],
    )
    note_box(
        doc,
        "Contraseña temporal",
        "El formato actual es Apellido.Nombre#Año. La persona debe cambiarla después del primer ingreso. El panel no conserva una copia visible de las contraseñas.",
        caution=True,
    )

    doc.add_heading("5. Perfil profesional", level=1)
    add_steps(
        doc,
        "Actualizar el perfil",
        [
            "Abra la sección Perfil.",
            "Edite nombre, cargo, institución, correo o biografía.",
            "Pulse Guardar perfil.",
            "Espere a que desaparezca el loader y verifique la vista pública.",
        ],
    )

    doc.add_heading("6. Cursos y archivos", level=1)
    add_steps(
        doc,
        "Crear un curso",
        [
            "Complete código, nombre, periodo, estado, facultad y descripción.",
            "Seleccione Actual o Histórico.",
            "Pulse Guardar curso.",
        ],
    )
    add_steps(
        doc,
        "Subir un archivo a un curso",
        [
            "Localice el curso.",
            "Seleccione la categoría del archivo.",
            "Defina la visibilidad como Público o Privado.",
            "Pulse Elegir archivo y seleccione el documento.",
            "Espere la confirmación antes de cerrar la página.",
        ],
    )
    add_bullets(
        doc,
        [
            "Público: puede ser consultado por cualquier visitante.",
            "Privado: solo es accesible desde la administración.",
            "Eliminar un curso también elimina sus relaciones y debe hacerse con precaución.",
        ],
    )

    doc.add_heading("7. Foros", level=1)
    add_steps(
        doc,
        "Crear un foro",
        [
            "Seleccione el curso.",
            "Escriba título y descripción.",
            "Defina el estado: Abierto, Cerrado o Solo lectura.",
            "Pulse Nuevo foro.",
        ],
    )
    add_steps(
        doc,
        "Responder como docente",
        [
            "Localice el foro.",
            "Seleccione si responderá al foro o a un mensaje.",
            "Escriba la respuesta y pulse Publicar respuesta.",
        ],
    )
    add_bullets(
        doc,
        [
            "Abierto: permite nuevas respuestas.",
            "Cerrado: conserva la conversación sin nuevas participaciones.",
            "Solo lectura: mantiene el contenido como referencia histórica.",
        ],
    )

    doc.add_heading("8. Publicaciones y trayectoria", level=1)
    add_steps(
        doc,
        "Agregar una publicación",
        [
            "Complete título, año, tipo, revista o evento, autores y estado.",
            "Pulse Agregar.",
            "Use el buscador para comprobar el registro.",
        ],
    )
    add_steps(
        doc,
        "Agregar un hito",
        [
            "Abra Línea del tiempo.",
            "Complete año, título, categoría y descripción.",
            "Pulse Nuevo hito.",
        ],
    )

    doc.add_heading("9. Contador de visitas", level=1)
    doc.add_paragraph(
        "El total se muestra discretamente en el pie de la página pública. Se registra una visita por sesión del navegador. El contador no almacena datos personales."
    )

    doc.add_heading("10. Configuración del registro", level=1)
    add_bullets(
        doc,
        [
            "NEXT_PUBLIC_ALLOW_PUBLIC_SIGNUP=false: las cuentas se crean desde el archivo Excel o CSV.",
            "NEXT_PUBLIC_ALLOW_PUBLIC_SIGNUP=true: el sitio muestra la opción de registro público.",
            "En producción el valor actual es false y Supabase también tiene deshabilitado el registro abierto.",
        ],
    )

    doc.add_heading("11. Seguridad y mantenimiento", level=1)
    add_bullets(
        doc,
        [
            "No compartir la contraseña administrativa.",
            "No publicar la clave SUPABASE_SERVICE_ROLE_KEY ni colocarla en archivos del repositorio.",
            "Descargar las credenciales temporales solo en un equipo seguro.",
            "Revisar cuidadosamente antes de eliminar cursos, archivos o publicaciones.",
            "Cerrar sesión al terminar, especialmente en equipos compartidos.",
            "Mantener actualizados los datos institucionales y enlaces profesionales.",
        ],
    )
    note_box(
        doc,
        "Loader",
        "Durante una edición o carga espere a que el indicador desaparezca. No recargue ni cierre la página mientras la operación está en curso.",
    )

    doc.add_heading("12. Verificación después de cambios", level=1)
    checklist = [
        "Abrir la vista pública.",
        "Comprobar el contenido en computadora y teléfono.",
        "Confirmar que los archivos públicos abren correctamente.",
        "Verificar que una cuenta normal no acceda al panel.",
        "Cerrar la sesión administrativa.",
    ]
    add_bullets(doc, checklist)

    path = OUTPUT / "Manual_de_administrador_Carlos_Lopez.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for manual in (build_user_manual(), build_admin_manual()):
        print(manual)
